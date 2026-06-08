"""Build a clean, ATS-safe .docx from a structured resume (no tables/columns/graphics, which ATS
parsers choke on). Uses python-docx, already a dependency. Also a deterministic ATS health score."""
import io


def build_docx(r: dict) -> bytes:
    import docx
    from docx.shared import Pt

    doc = docx.Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def heading(t):
        p = doc.add_paragraph()
        run = p.add_run(t.upper())
        run.bold = True
        run.font.size = Pt(12)

    name = (r.get("name") or "").strip()
    if name:
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(16)
    contact = " | ".join(x for x in [r.get("email"), r.get("phone"), *(r.get("links") or [])] if x)
    if contact:
        doc.add_paragraph(contact)

    if r.get("summary"):
        heading("Summary")
        doc.add_paragraph(r["summary"])

    if r.get("skills"):
        heading("Skills")
        doc.add_paragraph(", ".join(r["skills"]))

    if r.get("experience"):
        heading("Experience")
        for e in r["experience"]:
            line = " | ".join(x for x in [e.get("title"), e.get("company"), e.get("dates")] if x)
            if line:
                p = doc.add_paragraph()
                p.add_run(line).bold = True
            for b in (e.get("bullets") or []):
                doc.add_paragraph(b, style="List Bullet")

    if r.get("projects"):
        heading("Projects")
        for pr in r["projects"]:
            line = " | ".join(x for x in [pr.get("name"), pr.get("stack"), pr.get("dates")] if x)
            if line:
                p = doc.add_paragraph()
                p.add_run(line).bold = True
            for b in (pr.get("bullets") or []):
                doc.add_paragraph(b, style="List Bullet")

    for sec in (r.get("sections") or []):
        if sec.get("heading"):
            heading(sec["heading"])
        for it in (sec.get("items") or []):
            doc.add_paragraph(it, style="List Bullet")

    if r.get("education"):
        heading("Education")
        for ed in r["education"]:
            line = " | ".join(x for x in [ed.get("degree"), ed.get("school"), ed.get("dates")] if x)
            if line:
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Strong action verbs and weak openers, for the deterministic ATS health checks.
_ACTION = ("led", "built", "shipped", "designed", "architected", "implemented", "launched",
           "improved", "reduced", "increased", "drove", "owned", "delivered", "created", "scaled",
           "automated", "developed", "optimized", "migrated", "established")


def ats_health(r: dict) -> dict:
    """Deterministic resume-quality score (0-100) + concrete checks. Independent of any JD.
    Delegates to the ResumeWorded-style analyzer (Impact / Brevity / Style / Sections) so the score
    is explainable; keeps `score`+`checks` keys for existing callers and adds `categories`."""
    from .ats_rules import quality_report
    return quality_report(r)
