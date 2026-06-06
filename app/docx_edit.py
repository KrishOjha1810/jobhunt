"""Edit the user's ORIGINAL resume in place so exports keep their exact formatting.

We store the user's resume as a .docx (converting from PDF when needed), then apply only the changes
they accepted (summary rewrite, added skills, bullet rewrites) by matching paragraph text , every
other paragraph, style, and section in their document is left untouched. PDF->DOCX conversion is
lossy (best-effort); .docx originals are ideal.
"""
import base64
import io


def to_docx_b64(path):
    """Return base64 of a .docx for the uploaded resume at `path`. Converts PDF -> DOCX best-effort.
    Returns None when we can't produce a docx (caller falls back to the generated template)."""
    p = (path or "").lower()
    try:
        if p.endswith(".docx"):
            with open(path, "rb") as f:
                return base64.b64encode(f.read()).decode()
        if p.endswith(".pdf"):
            from pdf2docx import Converter
            import os
            out = path + ".converted.docx"
            cv = Converter(path)
            cv.convert(out)  # all pages
            cv.close()
            with open(out, "rb") as f:
                data = f.read()
            try:
                os.remove(out)
            except Exception:
                pass
            return base64.b64encode(data).decode()
    except Exception:
        return None
    return None


def _norm(s):
    return " ".join((s or "").split()).strip().lower()


def _set_text(p, text):
    """Replace a paragraph's text while keeping its paragraph/run style (first run keeps formatting)."""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def apply_edits(b64, edits):
    """Apply accepted edits to the stored docx; return new docx bytes (or None on failure).
    edits = {summary_old, summary_new, bullets:[{old,new}], skills_add:[...]}."""
    if not b64:
        return None
    try:
        from docx import Document
        doc = Document(io.BytesIO(base64.b64decode(b64)))
    except Exception:
        return None
    sum_old = _norm(edits.get("summary_old"))
    sum_new = edits.get("summary_new")
    bullets = { _norm(b.get("old")): b.get("new") for b in (edits.get("bullets") or [])
                if b.get("old") and b.get("new") }
    skills_add = [s for s in (edits.get("skills_add") or []) if s]
    skills_done = False
    for p in doc.paragraphs:
        raw = p.text
        t = _norm(raw)
        if not t:
            continue
        if sum_new and sum_old and t == sum_old:
            _set_text(p, sum_new); continue
        if t in bullets and bullets[t]:
            _set_text(p, bullets[t]); continue
        if skills_add and not skills_done and ("skill" in t or ("," in raw and len(raw.split(",")) >= 4)):
            add = [s for s in skills_add if s.lower() not in t]
            if add:
                _set_text(p, raw.rstrip(" .;,") + ", " + ", ".join(add))
                skills_done = True
    try:
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception:
        return None
